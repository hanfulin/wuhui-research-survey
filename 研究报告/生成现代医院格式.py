from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Inches


BASE = Path(__file__).resolve().parent
SRC = BASE / "文献综述初稿.md"
OUT = BASE / "文献综述初稿-现代医院格式.docx"


def set_run_font(run, size=None, bold=None, italic=None, east="宋体", west="Times New Roman"):
    run.font.name = west
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run._element.rPr.rFonts.set(qn("w:ascii"), west)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), west)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_format(p, size=9, first_line=False, align=None, before=0, after=0, line=13):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = Pt(line)
    if first_line:
        fmt.first_line_indent = Cm(0.55)
    if align is not None:
        p.alignment = align
    for run in p.runs:
        set_run_font(run, size=size)


def add_rich_text(p, text, size=9, bold_default=False, east="宋体", west="Times New Roman"):
    # Handles simple Markdown bold spans while preserving the original text.
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        clean = part[2:-2] if is_bold else part
        run = p.add_run(clean)
        set_run_font(run, size=size, bold=(bold_default or is_bold), east=east, west=west)


def set_cell_text(cell, text, size=8.2, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(11)
    add_rich_text(p, text.strip(), size=size, bold_default=bold)


def shade_cell(cell, fill="F2F2F2"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "808080")


def set_columns(section, num=2, space_twips=425):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols = cols[0] if cols else OxmlElement("w:cols")
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))
    if not cols.getparent():
        sect_pr.append(cols)


def make_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(9)
    normal.paragraph_format.line_spacing = Pt(13)
    normal.paragraph_format.space_after = Pt(0)

    for name, size in [("Heading 1", 10), ("Heading 2", 9.5), ("Heading 3", 9)]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st.font.size = Pt(size)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(5)
        st.paragraph_format.space_after = Pt(2)
        st.paragraph_format.line_spacing = Pt(12.5)


def configure_section(section):
    section.page_width = Pt(501.48)
    section.page_height = Pt(758.88)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.55)


def split_table_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def add_table(doc, lines):
    rows = [split_table_row(l) for l in lines if l.strip().startswith("|")]
    rows = [r for i, r in enumerate(rows) if i != 1]
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(table.cell(i, j), value, bold=(i == 0), align=align)
            if i == 0:
                shade_cell(table.cell(i, j))
    set_table_borders(table)
    p = doc.add_paragraph()
    set_para_format(p, size=8.2, after=3, line=10)


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = "现代医院  投稿排版稿"
    set_para_format(p, size=8, line=9)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_para_format(p, size=8, line=9)


def build():
    text = SRC.read_text(encoding="utf-8").replace("\r\n", "\n")
    lines = text.split("\n")

    doc = Document()
    make_styles(doc)
    configure_section(doc.sections[0])
    add_header_footer(doc.sections[0])

    in_body_columns = False
    table_buf = []
    prev_blank = False

    for raw in lines:
        line = raw.rstrip()

        if table_buf and (not line.startswith("|")):
            add_table(doc, table_buf)
            table_buf = []

        if not line.strip():
            prev_blank = True
            continue
        if line.strip() == "---":
            continue
        if line.startswith("|"):
            table_buf.append(line)
            continue

        h = re.match(r"^(#{1,4})\s+(.+)$", line)
        if h:
            level = len(h.group(1))
            title = h.group(2).strip()
            if title.startswith("1 ") and not in_body_columns:
                section = doc.add_section(WD_SECTION.CONTINUOUS)
                configure_section(section)
                set_columns(section, 2)
                in_body_columns = True
            if level == 1:
                p = doc.add_paragraph(title)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, size=15, bold=True, east="黑体")
                p.paragraph_format.space_after = Pt(5)
                p.paragraph_format.line_spacing = Pt(18)
            elif level == 2:
                # 摘要和 Abstract 保持居中；正文二级标题按期刊样式左对齐。
                p = doc.add_paragraph(title, style="Heading 1")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if not in_body_columns else WD_ALIGN_PARAGRAPH.LEFT
            elif level == 3:
                p = doc.add_paragraph(title, style="Heading 2")
            else:
                p = doc.add_paragraph(title, style="Heading 3")
            prev_blank = False
            continue

        # Front matter before the Chinese title.
        if not in_body_columns and line.startswith("**") and "：" in line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_rich_text(p, line, size=9)
            set_para_format(p, size=9, after=1, line=12)
            prev_blank = False
            continue

        if line.startswith("[") and re.match(r"^\[\d+\]", line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.45)
            p.paragraph_format.first_line_indent = Cm(-0.45)
            add_rich_text(p, line.strip(), size=8.2)
            set_para_format(p, size=8.2, line=11.5)
            continue

        p = doc.add_paragraph()
        is_abstract_para = (not in_body_columns and ("Objective " in line or line.startswith("**目的**") or line.startswith("**关键词**") or line.startswith("**Keywords**")))
        add_rich_text(p, line.strip(), size=8.7 if is_abstract_para else 9)
        set_para_format(
            p,
            size=8.7 if is_abstract_para else 9,
            first_line=in_body_columns,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            before=0,
            after=1 if prev_blank else 0,
            line=12.5 if in_body_columns else 12,
        )
        prev_blank = False

    if table_buf:
        add_table(doc, table_buf)

    # Privacy-oriented cleanup of document properties.
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.title = "公立医院科研经费包干制落地的内控视角综述"

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
